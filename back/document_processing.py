import asyncio
import os
import re
from typing import List, Optional

import PyPDF2
import docx

from utils import allowed_file


def split_text_into_chunks(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    if not text:
        return []

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
        if start >= len(text):
            if end < len(text):
                chunks.append(text[start:])
            break

    return [c.strip() for c in chunks if c.strip()]


def extract_text_from_pdf(file_path: str) -> Optional[str]:
    text = ""
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            if reader.is_encrypted:
                print(f"PDF '{os.path.basename(file_path)}' is encrypted.")

            for page in reader.pages:
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                except Exception as page_err:
                    print(f"Error extracting text from page: {page_err}")

            return text.strip() if text else None
    except Exception as e:
        print(f"Error reading PDF '{os.path.basename(file_path)}': {e}")
        return None


def extract_content_from_docx(file_path: str) -> Optional[str]:
    try:
        doc = docx.Document(file_path)
        full_content = []

        for element in doc.element.body:
            if isinstance(element, docx.oxml.text.paragraph.CT_P):
                para = docx.text.paragraph.Paragraph(element, doc)
                if para.text and para.text.strip():
                    full_content.append(para.text.strip())
            elif isinstance(element, docx.oxml.table.CT_Tbl):
                table = docx.table.Table(element, doc)
                if not table.rows:
                    continue

                markdown_table_lines = []
                num_cols = len(table.columns)

                try:
                    header_cells = []
                    for cell in table.rows[0].cells:
                        cell_text = "\n".join([p.text for p in cell.paragraphs]).strip()
                        cell_text_cleaned = re.sub(r'\s+', ' ', cell_text).strip().replace('|', '\\|')
                        header_cells.append(cell_text_cleaned if cell_text_cleaned else " ")

                    markdown_table_lines.append("| " + " | ".join(header_cells) + " |")
                    markdown_table_lines.append("|" + "---|" * num_cols)

                    for i, row in enumerate(table.rows):
                        if i == 0:
                            continue

                        row_cells_text = []
                        for cell in row.cells:
                            cell_text = "\n".join([p.text for p in cell.paragraphs]).strip()
                            cell_text_cleaned = re.sub(r'\s+', ' ', cell_text).strip().replace('|', '\\|')
                            row_cells_text.append(cell_text_cleaned if cell_text_cleaned else " ")

                        if len(row_cells_text) == num_cols:
                            markdown_table_lines.append("| " + " | ".join(row_cells_text) + " |")

                    full_content.append("\n```markdown\n" + "\n".join(markdown_table_lines) + "\n```\n")

                except Exception as table_err:
                    print(f"Error processing table: {table_err}")

        content = '\n\n'.join(full_content).strip()
        content = re.sub(r'\n{3,}', '\n\n', content)
        return content if content else None
    except Exception as e:
        print(f"Error extracting content from DOCX: {e}")
        return None


async def process_document(file_path: str, original_filename: str) -> Optional[List[str]]:
    if not allowed_file(original_filename):
        return None

    try:
        file_ext = original_filename.lower().rsplit('.', 1)[1]
        text = None

        if file_ext == 'txt':
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                text = await asyncio.to_thread(f.read)
        elif file_ext == 'pdf':
            text = await asyncio.to_thread(extract_text_from_pdf, file_path)
        elif file_ext == 'docx':
            text = await asyncio.to_thread(extract_content_from_docx, file_path)

        if not text:
            return None

        chunks = await asyncio.to_thread(split_text_into_chunks, text)
        return chunks
    except Exception as e:
        print(f"Error processing document: {e}")
        return None
